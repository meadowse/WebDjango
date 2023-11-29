from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Document_compose
import shutil
def createLabels(data):
    print(data)
    # Откройте файл с шаблоном
    template_doc = DocxTemplate("tes.docx")
    dock = Document()
    dock.save("generated_nakleiki.docx")
    # Создайте список с данными (имена и адреса)
    # Создайте новый документ для сохранения наклеек
    new_doc = Document()
    print("dd")
    # Создайте и добавьте наклейки на основе данных
    i = 1
    for item in data:
        if i % 2 != 0:
            context = {'name': item['name'], 'address': item['address'], 'number':i, 'index':item['index']}
        else :
            context = {'name1': item['name'], 'address1': item['address'], 'number1': i, 'index1': item['index']}
        template_doc.render(context)
        template_doc.save("temp.docx")  # Сохраните шаблон с данными во временный файл
        # Откройте временный файл и скопируйте его содержимое в новый документ
        # temp_doc = Document("temp.docx")
        dock =  Document_compose("generated_nakleiki.docx") # куда хотим сохранить
        composer = Composer(dock)
        doc2 = Document_compose("temp.docx")
        composer.append(doc2)
        composer.save("generated_nakleiki.docx")
        i+=1
    # Сохраните новый документ с наклейками
# new_doc.save("generated_nakleiki.docx")