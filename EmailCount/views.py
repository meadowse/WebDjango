# import mimetypes
# import os
import glob
import zipfile
from time import perf_counter
from xml.dom.minidom import Document
from docx.shared import Inches, Mm, Pt
# from docxcompose.composer import Composer
# from docx import Document as Document_compose
# from django.shortcuts import render
from django.http import HttpResponseForbidden, HttpResponse  # , HttpResponseNotFound
import SendMail
# from EmailCount.models import Posttable
from test import *
from Print_Stikers import *
from dbServer.views import *
# from django.http import JsonResponse
# import json
# import io
# from django.http import FileResponse
# from django.core.files.base import ContentFile
from django.views.decorators.csrf import csrf_exempt
from databaseConnect import *


def createDirs(user):
    if not os.path.exists("/home/meadowse/tmp/" + user + "/tmp"):
        os.makedirs("/home/meadowse/tmp/" + user + "/tmp")


# Create your views here.
@csrf_exempt
def emailmainCourt(request):
    # show = Posttable.objects.filter(MailingType=2).values()
    # data = []
    info = query_db("select * from private.universal_mailing u where u.type_mailing = 'Email'")
    # print(info)
    names = ["Наименование рассылки", "Шаблон", "Кол-во писем", "Дата создания", "Дата отправки", "Ответственный",
             "Статус", "Привязанные контакты"]
    Companies = query_db("select * from private.all_data_companies limit 10")
    newsleter = (query_db("select * from private.templates"))
    return render(request, 'EmailCount/mail.html', {'title': "main", 'mail': info, 'names': names,
                                                    'Companies': Companies, 'templateForSend': newsleter,
                                                    'types_mailing': '1'})


@csrf_exempt
def SendEmail(request):
    obj = json.loads(request.body)
    recipients = []
    for i in obj:
        tmp = i[0].split(",")
        print(tmp)
        if tmp[0] != 'on':
            Info = (query_db("select mc.id_company from private.mailing_companies mc where id_mailing = %s", (tmp[0],)))
            # TODO получить информацию фио почту и остальное о человеке
            for companyId in Info:
                recipients.append(query_db("select legal_address, mail, fio_head  from private.companies where id = %s",
                                           (str(companyId['id_company']),)))
                print(recipients)
            # recipients.append(Info.mail)
            # recipients.append(Info.name)
            # recipients.append(Info.addres)
        # записываем в форму информацию о людях
    err = 0
    if recipients:
        err = SendMail.sendMail(recipients)
    return HttpResponse(f"<h4>HelloWorld</h4><p>{{err}}</p>")


@csrf_exempt
def addTarget(request):
    print(request.GET)
    obj = json.loads(request.body)
    print(obj)
    id_company = []
    j = 0
    for i in obj:
        print(int(i['check_' + str(j)][0].split(",")[0]))
        id_company.append(((int(i['check_' + str(j)][0].split(",")[0])), int(i['check_' + str(j)][0].split(",")[1])))
        #  id_company.append()
        j += 1
    query = "insert into private.mailing_companies values " + ",".join("(%s, %s)" for i in id_company)
    flattened_values = [item for sublist in id_company for item in sublist]
    Insert_db(query, flattened_values)
    return HttpResponse(f"<h4>HelloWorld</h4><p>{{err}}</p>")


@csrf_exempt
def removeTarget(request):
    obj = json.loads(request.body)
    j = 0
    id_company = []
    type_mailing = ''
    for i in obj:
        id_company.append(int(i['check_'+str(j)][0].split(",")[0]))
        type_mailing = i['check_'+str(j)][0].split(",")[1]
        j += 1
    print(id_company)
    print(type_mailing)
    query = ("DELETE FROM private.mailing_companies "
             "WHERE id_company IN ({}) AND id_mailing = %s").format(','.join(['%s'] * len(id_company)))
    Insert_db(query, id_company + [type_mailing])
    return HttpResponse(f"<h4>HelloWorld</h4><p>{{err}}</p>")


def createTempFile(template_doc, context, composer, user):
    template_doc.render(context)
    if not os.path.isfile(os.path.join("/home/meadowse/tmp/" + user + "/tmp/", 'temp.docx')):
        document = Document()
        fileName = "/home/meadowse/tmp/" + user + "/tmp/temp.docx"
        linesInDocx(document, fileName)
    template_doc.save("/home/meadowse/tmp/" + user + "/tmp/temp.docx")
    # Сохраните шаблон с данными во временный файл
    doc2 = Document_compose("/home/meadowse/tmp/" + user + "/tmp/temp.docx")
    composer.append(doc2)


# настройка отступов
def linesInDocx(document, file_name, top=0, down=0, left=0.0, right=0.0, height=297, width=210):
    section = document.sections[-1]
    style = document.styles['Normal']
    font = style.font
    font.name = 'Montserrat'
    font.size = Pt(10)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(down)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.page_height = Mm(height)
    section.page_width = Mm(width)
    document.save(file_name)


# передаем имя файла для скачивания
def printDocx(excel_file_name):
    with open(excel_file_name, "rb") as fp:
        response = HttpResponse(fp.read())
    file_type = mimetypes.guess_type(excel_file_name)
    if file_type is None:
        file_type = 'application/octet-stream'
    response['Content-Type'] = file_type
    response['Content-Length'] = str(os.stat(excel_file_name).st_size)
    response['Content-Disposition'] = "attachment; filename=excel_file.xlsx"
    return response


@csrf_exempt
def createLabels(request):
    start = perf_counter()
    user = request.user.username
    folder_path = '/home/meadowse/tmp/' + user  # Замените на путь к вашей папке
    # Создайте временный zip-архив
    zip_file_path = os.path.join(folder_path + "/tmp/", 'temp.zip')
    files = glob.glob('/home/meadowse/tmp/' + user + "/*.docx")
    for f in files:
        os.remove(f)
    if os.path.isfile(zip_file_path):
        os.remove(zip_file_path)
    createDirs(user)
    obj = json.loads(request.body)
    Info = []
    for i in obj:
        tmp = i[0].split(",")
        if tmp[0] != 'on':
            Info = (query_db("SELECT id, name, legal_address, position_head_dp, fio_dp, index FROM companies c "
                             "join private.mailing_companies m on c.id = m.id_company WHERE id_mailing = %s "
                             "order by id", (tmp[0],)))  # TODO получить информацию фио почту и остальное о человеке
    # Откройте файл с шаблоном
    template_doc = DocxTemplate("tes.docx")
    # Создайте и добавьте наклейки на основе данных
    fileName = ''
    dock = Document_compose()
    composer = Composer(dock)
    for i in range(0, len(Info), 2):
        context = {'name1': Info[i].get('position_head_dp') + ' ' + Info[i].get('name') + ' ' + Info[i].get('fio_dp'),
                   'address1': Info[i].get('legal_address'), 'number1': i + 1, 'index1': Info[i].get('index')}
        if i + 1 < len(Info):
            context.update({
                'name2': Info[i + 1].get('position_head_dp') + ' ' + Info[i + 1].get('name') + ' ' + Info[i + 1].get('fio_dp'),
                'address2': Info[i + 1].get('legal_address'), 'number2': i + 2, 'index2': Info[i + 1].get('index')})
        if i % 100 == 0:
            if fileName != '':
                composer.save(fileName)
            document = Document()
            fileName = "/home/meadowse/tmp/" + user + "/generated_nakleiki_" + str(i // 100) + ".docx"
            linesInDocx(document, fileName)
            dock = Document_compose(fileName)  # куда хотим сохранить
            composer = Composer(dock)
        createTempFile(template_doc, context, composer, user)  # склейка шаблонов
        print("Add", (i + 1))
    composer.save(fileName)
    # Создайте временный zip-архив
    with zipfile.ZipFile(zip_file_path, 'w') as zip_file:
        # Рекурсивно добавьте все файлы из папки в zip-архив
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if file == "temp.zip" or file == "temp.docx":
                    continue
                file_path = os.path.join(root, file)
                zip_file.write(file_path, os.path.relpath(file_path, folder_path))
    # Откройте zip-архив и считайте его содержимое
    with open(zip_file_path, 'rb') as zip_file:
        response = HttpResponse(zip_file.read(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="folder.zip"'
    end = perf_counter()
    print(end - start)
    return response


@csrf_exempt
def createLetters(request):
    start = perf_counter()
    user = request.user.username
    folder_path = '/home/meadowse/tmp/' + user  # Замените на путь к вашей папке
    # Создайте временный zip-архив
    zip_file_path = os.path.join(folder_path + "/tmp/", 'temp.zip')
    files = glob.glob('/home/meadowse/tmp/' + user + "/*.docx")
    for f in files:
        os.remove(f)
    if os.path.isfile(zip_file_path):
        os.remove(zip_file_path)
    createDirs(user)
    obj = json.loads(request.body)  # читаем тело запроса
    tmp = obj.split("; ")
    i = tmp[1]
    if tmp[0] == "createLetters":
        tmp2 = (query_db("SELECT id, name, position_head_dp, fio_dp, io, appeal FROM companies c "
                         "join private.mailing_companies m on c.id = m.id_company WHERE id_mailing = %s order by id",
                         (i,)))
        template_doc = DocxTemplate("templateForMail.docx")
        # todo заменить шаблолн в зависимоти от того что отправить надо
        top = 0
        down = 0
        left = 0
        right = 0
    elif tmp[0] == "commercial":
        tmp2 = (query_db("select distinct c.id, c.company_name as name, c.fio_head, co.objects_owned, "
                         "ARRAY[zos(co.square)] as zos, ARRAY[agr(co.square)] as agr, ARRAY[days(co.square)] as day, "
                         "ARRAY[(zos(co.square) + agr(co.square))] as total from private.companies c "
                         "join private.mailing_companies m on c.id = m.id_company and m.id_mailing = %s "
                         "left join (select c.id, "
                         "array_agg((cn.cadastral_number || ' ' || cn.address)) as objects_owned, "
                         "min(cn.square) as square from private.companies c "
                         "join private.mailing_companies m on c.id = m.id_company and m.id_mailing = %s "
                         "join private.rights r on c.id = r.id_company "
                         "join private.types_law t on r.id_type_right = t.id "
                         "and (t.type_right = 'Собственность' or t.type_right = 'Аренда') "
                         "join private.cadastral_numbers cn on r.id_cadastral_number = cn.id "
                         "and days(cn.square) is not null join private.types_objects tob on cn.id_object_type = tob.id "
                         "and tob.object_type != 'Земельный участок' group by c.id "
                         "order by c.id) as co on c.id = co.id order by c.id", (i, i,)))
        template_doc = DocxTemplate("leghal.docx")  # todo заменить шаблолн в зависимоти от того что отправить надо
        top = 1.25
        down = 1.25
        left = 0.5
        right = 0.127
    else:
        tmp2 = (query_db("select c.id, name, substr(pc.legal_address, 9, length(pc.legal_address)) as legal_address, "
                         "io, appeal, pc.fio_head, "
                         "split_part(m.mailing_list_name, ' ', 1) || ' ' || pc.inn_company as number, "
                         "m.date_creation::date::text as date from companies c "
                         "join private.companies pc on c.id = pc.id "
                         "join private.mailing_companies mc on mc.id_company = pc.id "
                         "join private.mailing m on m.id = mc.id_mailing where m.id = %s order by c.id", (i,)))
        template_doc = DocxTemplate("Info.docx")  # todo заменить шаблолн в зависимоти от того что отправить надо
        top = 0
        down = 0
        left = 0
        right = 0
    j = 0
    fileName = ''
    dock = Document_compose()
    composer = Composer(dock)
    for i in range(0, len(tmp2)):
        if len(tmp2[i]) > 0:
            if i % 100 == 0:
                if fileName != '':
                    composer.save(fileName)
                document = Document()
                fileName = "/home/meadowse/tmp/" + user + "/MailsForSend_" + str(i // 100) + ".docx"
                linesInDocx(document, fileName, top, down, left, right)
                dock = Document_compose(fileName)  # куда хотим сохранить
                composer = Composer(dock)
            tmp2[i]['i'] = i + 1
            if tmp[0] == "commercial":
                if tmp2[i].get('zos')[0] is not None:
                    tmp2[i].get('zos')[0] = "{:,d}".format(tmp2[i].get('zos')[0]).replace(',', ' ')
                if tmp2[i].get('agr')[0] is not None:
                    tmp2[i].get('agr')[0] = "{:,d}".format(tmp2[i].get('agr')[0]).replace(',', ' ')
                if tmp2[i].get('total')[0] is not None:
                    tmp2[i].get('total')[0] = "{:,d}".format(tmp2[i].get('total')[0]).replace(',', ' ')
                    createTempFile(template_doc, tmp2[i], composer, user)  # склейка шаблонов
            else:
                createTempFile(template_doc, tmp2[i], composer, user)  # склейка шаблонов
            print("Add", (i + 1))
        j += 1
    composer.save(fileName)
    # Создайте временный zip-архив
    with zipfile.ZipFile(zip_file_path, 'w') as zip_file:
        # Рекурсивно добавьте все файлы из папки в zip-архив
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if file == "temp.zip" or file == "temp.docx":
                    continue
                file_path = os.path.join(root, file)
                zip_file.write(file_path, os.path.relpath(file_path, folder_path))
    # Откройте zip-архив и считайте его содержимое
    with open(zip_file_path, 'rb') as zip_file:
        response = HttpResponse(zip_file.read(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="folder.zip"'
    end = perf_counter()
    print(end - start)
    return response


@csrf_exempt
def removeLine(request):
    try:
        obj = json.loads(request.body)
        j = 0
        id_company = []
        print(obj)
        for i in obj:
            id_company.append(int(i['check_' + str(j)][0].split(",")[0]))
            j += 1
        query = "DELETE FROM private.mailing WHERE id IN ({})".format(','.join(['%s'] * len(id_company)))
        Insert_db(query, id_company)
        return HttpResponse(status=200)
    except:
        return HttpResponse(status=500)


def EmalePajeinfo(request, catid):
    rez1 = ("select * from private.all_data_companies a where a.inn_company::bigint not in (select mc.id_company "
            "from private.mailing_companies mc where mc.id_mailing = %s)" % catid)
    rez2 = ("select * from private.all_data_companies a where a.inn_company::bigint in (select mc.id_company "
            "from private.mailing_companies mc where mc.id_mailing = %s)" % catid)
    Info = search(request, rez1)
    info2 = search(request, rez2)
    return render(request, 'sms/updateMailing.html', {'NotInNewsletter': Info, 'InNewsletter': info2,
                                                      'template': catid, 'types_mailing': '1'})


@csrf_exempt
def changeStatus(request):
    try:
        obj = json.loads(request.body)  # читаем тело запроса
        print(obj)
        # tmp = obj.split(";")
        Insert_db("Update private.mailing set status = 'Отправлено', date_dispatch = now() where id = %s", (obj,))
        return HttpResponse(status=200)
    except:
        return HttpResponse(status=500)
